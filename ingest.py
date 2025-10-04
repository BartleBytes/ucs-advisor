import logging
import os
from pathlib import Path

import pandas as pd
from dotenv import load_dotenv
from docx import Document as DocxDocument
from zipfile import BadZipFile
from pypdf import PdfReader

from data_loaders import load_course_catalog
from langchain_community.embeddings import OllamaEmbeddings
from langchain_community.vectorstores import Chroma
from langchain_core.documents import Document
from langchain_openai import OpenAIEmbeddings


COURSE_CATALOG_PATH = "sources/Fall_2005_Class_List/Fall 2025 Class List_ZW.xlsx"
SOURCE_DIR = Path("sources")
PLAN_DIR = SOURCE_DIR / "4-Year_Plans_2024-25"
DESC_DIR = SOURCE_DIR / "Business_Course_Catalog_Descriptions"
TEMPLATE_DIR = SOURCE_DIR / "Degree Plan Templates"
MINOR_DIR = SOURCE_DIR / "Minors"


def read_docx(path: Path) -> str:
    try:
        doc = DocxDocument(path)
    except BadZipFile:
        print(f"[warn] Skipping non-docx file: {path}")
        return ""
    return "\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())


def read_pdf(path: Path) -> str:
    reader = PdfReader(path)
    return "\n".join(page.extract_text() or "" for page in reader.pages)


logger = logging.getLogger(__name__)


def main():
    load_dotenv()
    logging.basicConfig(level=logging.INFO, format="%(levelname)s:%(name)s:%(message)s")
    embed_model = os.getenv("EMBED_MODEL", os.getenv("OPENAI_EMBED_MODEL", "text-embedding-3-small"))
    ollama_host = os.getenv("OLLAMA_HOST", "http://localhost:11434")
    backend = os.getenv("LLM_BACKEND", "ollama").lower()
    openai_key = os.getenv("OPENAI_API_KEY") or os.getenv("LLM_OPENAI_API_KEY")

    # Load structured course data
    course_df = load_course_catalog(COURSE_CATALOG_PATH)
    logger.info("Loaded %d rows from course catalog %s", len(course_df), COURSE_CATALOG_PATH)

    schedule_mask = (
        course_df["days"].astype(str).str.strip().ne("")
        & course_df["days"].astype(str).str.strip().ne("-")
        & course_df["start_time"].astype(str).str.strip().ne("")
        & course_df["end_time"].astype(str).str.strip().ne("")
    )
    skipped_rows = len(course_df) - int(schedule_mask.sum())
    if skipped_rows:
        logger.info("Skipping %d catalog rows with missing schedule info", skipped_rows)
    course_df = course_df[schedule_mask].reset_index(drop=True)

    docs: list[Document] = []
    course_docs = 0
    plan_docs = 0
    desc_docs = 0
    template_docs = 0
    minor_docs = 0

    for _, r in course_df.iterrows():
        text = (
            f"course_id: {r['course_id']}\n"
            f"title: {r['title']}\n"
            f"term: {r['term']}\n"
            f"days: {r['days']}\n"
            f"time: {r['start_time']}-{r['end_time']}\n"
            f"modality: {r['modality']}\n"
            f"campus: {r['campus']}\n"
            f"location: {r['location']} {r['room']}\n"
            f"instructor: {r['instructor']}\n"
            f"instructor_email: {r['instructor_email']}\n"
            f"section: {r['section']}\n"
            f"class_number: {r['class_number']}\n"
            f"component: {r['component']}\n"
            f"credits: {r['credits']}\n"
            f"prerequisites: {r['prerequisites']}\n"
        )
        docs.append(
            Document(
                page_content=text,
                metadata={
                    "course_id": r["course_id"],
                    "term": r["term"],
                    "section": r["section"],
                    "class_number": r["class_number"],
                },
            )
        )
        course_docs += 1

    logger.info("Structured course documents prepared: %d", course_docs)

    # 4-year plans (.docx)
    if PLAN_DIR.exists():
        for path in PLAN_DIR.glob("*.docx"):
            content = read_docx(path)
            if not content:
                continue
            docs.append(
                Document(
                    page_content=content,
                    metadata={
                        "source": str(path),
                        "category": "four-year-plan",
                        "plan": path.stem.replace("_4yp_2024-25", "").replace("_", " "),
                    },
                )
            )
            plan_docs += 1
        logger.info("Loaded %d four-year plan documents from %s", plan_docs, PLAN_DIR)
    else:
        logger.info("Plan directory missing: %s", PLAN_DIR)

    # Course descriptions (.docx)
    if DESC_DIR.exists():
        for path in DESC_DIR.glob("*.docx"):
            content = read_docx(path)
            if not content:
                continue
            docs.append(
                Document(
                    page_content=content,
                    metadata={"source": str(path), "category": "course-description"},
                )
            )
            desc_docs += 1
        logger.info("Loaded %d course description documents from %s", desc_docs, DESC_DIR)
    else:
        logger.info("Course description directory missing: %s", DESC_DIR)

    # Degree plan templates (.xlsx)
    if TEMPLATE_DIR.exists():
        for path in TEMPLATE_DIR.glob("*.xlsx"):
            if path.name.startswith("~$"):
                continue
            plan_df = pd.read_excel(path)
            docs.append(
                Document(
                    page_content=f"Degree Plan Template ({path.stem}):\n{plan_df.to_csv(index=False)}",
                    metadata={"source": str(path), "category": "degree-plan-template"},
                )
            )
            template_docs += 1
        logger.info("Loaded %d degree plan templates from %s", template_docs, TEMPLATE_DIR)
    else:
        logger.info("Template directory missing: %s", TEMPLATE_DIR)

    # Minor requirements (PDF)
    if MINOR_DIR.exists():
        for path in MINOR_DIR.glob("*.pdf"):
            docs.append(
                Document(
                    page_content=read_pdf(path),
                    metadata={"source": str(path), "category": "minor"},
                )
            )
            minor_docs += 1
        logger.info("Loaded %d minor requirement documents from %s", minor_docs, MINOR_DIR)
    else:
        logger.info("Minor directory missing: %s", MINOR_DIR)

    logger.info("Total documents assembled for embedding: %d", len(docs))

    # Build embeddings
    if backend == "openai":
        if not openai_key:
            raise RuntimeError("OPENAI_API_KEY is required when LLM_BACKEND=openai")
        emb = OpenAIEmbeddings(model=embed_model, api_key=openai_key)
    else:
        emb = OllamaEmbeddings(model=embed_model, base_url=ollama_host)

    persist_dir = os.getenv("CHROMA_DIR", "chroma_ucd")
    vs = Chroma.from_documents(docs, emb, persist_directory=persist_dir)
    vs.persist()
    logger.info("Index built at %s/ with %d documents", persist_dir, len(docs))


if __name__ == "__main__":
    main()
